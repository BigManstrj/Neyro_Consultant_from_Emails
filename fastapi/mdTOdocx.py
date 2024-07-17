from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class MdToDocxConverter:
    def __init__(self, md_file, docx_file):
        self.md_file = md_file
        self.docx_file = docx_file

    def convert(self):
        # Чтение файла Markdown
        with open(self.md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Создание нового документа DOCX
        document = Document()

        # Разбиваем содержимое на строки
        lines = md_content.split('\n')

        for line in lines:
            if line.startswith('# '):
                # Заголовок уровня 1
                self.add_heading(document, line[2:], level=1)
            elif line.startswith('## '):
                # Заголовок уровня 2
                self.add_heading(document, line[3:], level=2)
            elif line.startswith('### '):
                # Заголовок уровня 3
                self.add_heading(document, line[4:], level=3)
            else:
                # Обычный текст
                self.add_paragraph(document, line)

        # Сохранение документа DOCX
        document.save(self.docx_file)

    def add_heading(self, document, text, level):
        heading = document.add_heading(level=level)
        run = heading.add_run(text)
        run.font.size = Pt(14) if level == 1 else Pt(12)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def add_paragraph(self, document, text):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.font.size = Pt(11)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


class DocxToMdConverter:
    def __init__(self, docx_file, md_file):
        self.docx_file = docx_file
        self.md_file = md_file

    def convert(self):
        # Чтение файла DOCX
        document = Document(self.docx_file)

        # Открытие файла Markdown для записи
        with open(self.md_file, 'w', encoding='utf-8') as f:
            previous_line_empty = False
            for para in document.paragraphs:
                if para.style.name.startswith('Heading 1'):
                    # Заголовок уровня 1
                    f.write('# ' + para.text + '\n')
                    previous_line_empty = False
                elif para.style.name.startswith('Heading 2'):
                    # Заголовок уровня 2
                    f.write('## ' + para.text + '\n')
                    previous_line_empty = False
                elif para.style.name.startswith('Heading 3'):
                    # Заголовок уровня 3
                    f.write('### ' + para.text + '\n')
                    previous_line_empty = False
                else:
                    if para.text.strip() == '':
                        if not previous_line_empty:
                            f.write('\n')
                            previous_line_empty = True
                    else:
                        f.write(para.text + '\n')
                        previous_line_empty = False

    '''  конвератция с дублированием вопроса как простой текст'''
    def convert_w2q(self):
        # Чтение файла DOCX
        document = Document(self.docx_file)

        # Открытие файла Markdown для записи
        with open(self.md_file, 'w', encoding='utf-8') as f:
            previous_line_empty = False
            for para in document.paragraphs:
                if para.style.name.startswith('Heading 1'):
                    # Заголовок уровня 1
                    # f.write('# ' + para.text + '\n')
                    f.write('\n')  # заменим на перенос строки
                    previous_line_empty = False
                elif para.style.name.startswith('Heading 2'):
                    # Заголовок уровня 2
                    # f.write('## ' + para.text + '\n')
                    f.write('\n')   # заменим на перенос строки
                    previous_line_empty = False
                elif para.style.name.startswith('Heading 3'):
                    # Заголовок уровня 3
                    f.write('### ' + para.text + '\n')
                    f.write(para.text[8:] + '\n')   # продублируем вопрос как простой текст для лучшего поиска в базе
                    previous_line_empty = False
                else:
                    if para.text.strip() == '':
                        if not previous_line_empty:
                            f.write('\n')
                            previous_line_empty = True
                    else:
                        f.write(para.text + '\n')
                        previous_line_empty = False

def convert_md_to_docx(md_file, docx_file):
    converter = MdToDocxConverter(md_file, docx_file)
    converter.convert()


def convert_docx_to_md(docx_file, md_file):
    converter = DocxToMdConverter(docx_file, md_file)
    converter.convert()

def convert_docx_to_md_w2q(docx_file, md_file):
    converter = DocxToMdConverter(docx_file, md_file)
    converter.convert_w2q()

if __name__ == '__main__':
    import argparse

    parser = argparse.ArgumentParser(description='Convert between Markdown and DOCX files.')
    parser.add_argument('mode', choices=['md_to_docx', 'docx_to_md'], help='Conversion mode: md_to_docx or docx_to_md')
    parser.add_argument('input_file', help='Input file path')
    parser.add_argument('output_file', help='Output file path')

    args = parser.parse_args()

    if args.mode == 'md_to_docx':
        convert_md_to_docx(args.input_file, args.output_file)
    else:
        convert_docx_to_md(args.input_file, args.output_file)

    print(f'Conversion from {args.input_file} to {args.output_file} completed successfully.')
